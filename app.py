"""
╔══════════════════════════════════════════════════════════════════════════════════╗
║         AI RESEARCH PAPER ANALYZER — Industry-Standard RAG Pipeline            ║
║   Formats: PDF (text + scanned), JPG/JPEG/PNG, URL, TXT, DOCX, CSV, MD         ║
║   LLM: Groq LLaMA 3.3-70B | Embeddings: MiniLM-L6-v2 | Store: FAISS           ║
╚══════════════════════════════════════════════════════════════════════════════════╝
"""

# ─── Standard Library ────────────────────────────────────────────────────────
import os
import io
import re
import time
import hashlib
import textwrap
from pathlib import Path
from datetime import datetime

# Set USER_AGENT before any LangChain imports to silence warnings
os.environ.setdefault(
    "USER_AGENT",
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 Chrome/120.0.0.0 Safari/537.36"
)

# ─── Third-Party ─────────────────────────────────────────────────────────────
import streamlit as st
from dotenv import load_dotenv

# LangChain ecosystem
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain_core.documents import Document

# ─── Load Environment ─────────────────────────────────────────────────────────
load_dotenv()

# ══════════════════════════════════════════════════════════════════════════════
#  PAGE CONFIG & GLOBAL STYLES
# ══════════════════════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="AI Research Analyzer",
    page_icon="🔬",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        "Get Help": "https://github.com",
        "About": "AI Research Paper Analyzer | Built with LangChain + Groq + FAISS"
    }
)

# ─── Dark Professional CSS ───────────────────────────────────────────────────
st.markdown("""
<style>
/* ── Import Google Font ── */
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');

/* ── Root Overrides ── */
html, body, [class*="css"] { font-family: 'Inter', sans-serif; }

/* ── Main Background ── */
.stApp { background: linear-gradient(135deg, #0d1117 0%, #161b22 50%, #0d1117 100%); }

/* ── Sidebar ── */
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #161b22 0%, #0d1117 100%);
    border-right: 1px solid #21262d;
}
[data-testid="stSidebar"] .stMarkdown h2,
[data-testid="stSidebar"] .stMarkdown h3 {
    color: #58a6ff;
    font-weight: 600;
}

/* ── Header ── */
.app-header {
    background: linear-gradient(90deg, #1f6feb 0%, #388bfd 50%, #58a6ff 100%);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    background-clip: text;
    font-size: 2.4rem;
    font-weight: 700;
    text-align: center;
    padding: 0.5rem 0;
    letter-spacing: -0.5px;
}
.sub-header {
    text-align: center;
    color: #8b949e;
    font-size: 1rem;
    margin-top: -0.6rem;
    margin-bottom: 1.5rem;
}

/* ── Cards ── */
.analysis-card {
    background: #161b22;
    border: 1px solid #21262d;
    border-radius: 12px;
    padding: 1.2rem 1.5rem;
    margin: 0.8rem 0;
    transition: border-color 0.2s;
}
.analysis-card:hover { border-color: #388bfd; }

/* ── Section labels ── */
.section-label {
    color: #58a6ff;
    font-size: 0.78rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 1.2px;
    margin-bottom: 0.3rem;
}
.section-content {
    color: #c9d1d9;
    font-size: 0.95rem;
    line-height: 1.7;
}

/* ── Badges ── */
.badge {
    display: inline-block;
    padding: 0.15rem 0.6rem;
    border-radius: 20px;
    font-size: 0.72rem;
    font-weight: 600;
    margin: 0.2rem;
}
.badge-blue  { background: #1f3a5f; color: #58a6ff; border: 1px solid #1f6feb; }
.badge-green { background: #1a3a2a; color: #56d364; border: 1px solid #2ea043; }
.badge-amber { background: #3a2a0a; color: #e3b341; border: 1px solid #9e6a03; }
.badge-red   { background: #3a1a1a; color: #f85149; border: 1px solid #da3633; }

/* ── Source box ── */
.source-box {
    background: #0d1117;
    border: 1px solid #30363d;
    border-left: 3px solid #388bfd;
    border-radius: 6px;
    padding: 0.7rem 1rem;
    margin: 0.4rem 0;
    font-size: 0.82rem;
    color: #8b949e;
    font-family: 'Courier New', monospace;
}

/* ── Chat bubbles ── */
.chat-user {
    background: #1f3a5f;
    border-radius: 12px 12px 4px 12px;
    padding: 0.8rem 1.1rem;
    margin: 0.5rem 0;
    color: #c9d1d9;
    border: 1px solid #1f6feb;
}
.chat-assistant {
    background: #161b22;
    border-radius: 12px 12px 12px 4px;
    padding: 0.8rem 1.1rem;
    margin: 0.5rem 0;
    color: #c9d1d9;
    border: 1px solid #21262d;
}

/* ── Metric boxes ── */
.metric-box {
    background: #161b22;
    border: 1px solid #21262d;
    border-radius: 8px;
    padding: 1rem;
    text-align: center;
}
.metric-value { color: #58a6ff; font-size: 1.6rem; font-weight: 700; }
.metric-label { color: #8b949e; font-size: 0.78rem; margin-top: 0.2rem; }

/* ── Streamlit element overrides ── */
.stTextInput input, .stTextArea textarea {
    background: #0d1117 !important;
    border: 1px solid #30363d !important;
    color: #c9d1d9 !important;
    border-radius: 8px !important;
}
.stButton > button {
    background: linear-gradient(90deg, #1f6feb, #388bfd);
    color: white;
    border: none;
    border-radius: 8px;
    font-weight: 600;
    font-size: 0.9rem;
    padding: 0.5rem 1.2rem;
    transition: all 0.2s;
    width: 100%;
}
.stButton > button:hover {
    background: linear-gradient(90deg, #388bfd, #58a6ff);
    transform: translateY(-1px);
    box-shadow: 0 4px 12px rgba(56,139,253,0.3);
}
.stTabs [data-baseweb="tab"] {
    background: transparent;
    color: #8b949e;
    border-bottom: 2px solid transparent;
    font-weight: 500;
}
.stTabs [aria-selected="true"] {
    color: #58a6ff !important;
    border-bottom: 2px solid #388bfd !important;
}
.stAlert { border-radius: 8px; }

div[data-testid="stExpander"] {
    background: #161b22;
    border: 1px solid #21262d;
    border-radius: 8px;
}

/* ── Progress steps ── */
.step-row {
    display: flex;
    align-items: center;
    gap: 0.7rem;
    padding: 0.4rem 0;
    color: #c9d1d9;
    font-size: 0.9rem;
}
.step-icon { font-size: 1.1rem; min-width: 1.5rem; }
</style>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
#  SESSION STATE INITIALIZATION
# ══════════════════════════════════════════════════════════════════════════════
def init_session():
    defaults = {
        "vector_store": None,
        "doc_hash": None,
        "doc_metadata": {},
        "chat_history": [],
        "raw_text": "",
        "processing_done": False,
        "doc_chunks": [],
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

init_session()


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT LOADERS
# ══════════════════════════════════════════════════════════════════════════════

def _file_hash(content: bytes) -> str:
    """SHA-256 fingerprint to detect document changes → avoids reprocessing."""
    return hashlib.sha256(content).hexdigest()


def load_pdf(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """
    Two-pass PDF loader:
    Pass 1 → pypdf (fast, text-based PDFs)
    Pass 2 → pdfplumber (better layout extraction for dense papers)
    Pass 3 → OCR via pdf2image + pytesseract (scanned/image PDFs)
    """
    text = ""
    meta = {"source": filename, "type": "pdf", "pages": 0}
    
    # ── Pass 1: pypdf ──
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        meta["pages"] = len(reader.pages)
        for page in reader.pages:
            t = page.extract_text() or ""
            text += t + "\n"
        text = text.strip()
    except Exception:
        pass

    # ── Pass 2: pdfplumber (if pypdf got little text) ──
    if len(text) < 500:
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                meta["pages"] = len(pdf.pages)
                text = ""
                for page in pdf.pages:
                    t = page.extract_text() or ""
                    text += t + "\n"
            text = text.strip()
        except ImportError:
            pass
        except Exception:
            pass

    # ── Pass 3: OCR (for scanned PDFs) ──
    if len(text) < 300:
        text = _ocr_pdf(file_bytes, meta)

    if not text.strip():
        raise ValueError(
            "❌ Could not extract text from this PDF.\n"
            "• Text-layer extraction failed\n"
            "• pdfplumber extraction failed\n"
            "• OCR extraction failed\n\n"
            "Try: Ensure the PDF is not password-protected."
        )
    
    meta["char_count"] = len(text)
    return text, meta


def _ocr_pdf(file_bytes: bytes, meta: dict) -> str:
    """OCR fallback for scanned PDFs using pdf2image + pytesseract."""
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
        images = convert_from_bytes(file_bytes, dpi=200)
        meta["pages"] = len(images)
        meta["ocr_used"] = True
        texts = []
        for img in images:
            t = pytesseract.image_to_string(img, lang="eng")
            texts.append(t)
        return "\n\n".join(texts)
    except ImportError:
        return ""
    except Exception:
        return ""


def load_image(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """
    OCR loader for JPG, JPEG, PNG images using pytesseract.
    Falls back to a helpful error if Tesseract is not installed.
    """
    meta = {"source": filename, "type": "image"}
    try:
        from PIL import Image
        import pytesseract
        img = Image.open(io.BytesIO(file_bytes))
        # Pre-process: convert to RGB, scale up for better OCR
        img = img.convert("RGB")
        # Scale up small images for better OCR accuracy
        w, h = img.size
        if w < 1000:
            scale = 1000 / w
            img = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        text = pytesseract.image_to_string(img, lang="eng")
        meta["char_count"] = len(text)
        meta["image_size"] = f"{w}x{h}"
        if not text.strip():
            raise ValueError("OCR returned empty text. The image may not contain readable text.")
        return text, meta
    except ImportError as e:
        if "pytesseract" in str(e):
            raise RuntimeError(
                "🔧 **pytesseract not installed.**\n"
                "Run: `pip install pytesseract pillow`\n"
                "Also install Tesseract OCR binary from: https://github.com/UB-Mannheim/tesseract/wiki"
            )
        raise RuntimeError(f"Image loading failed: {e}")


def load_url(url: str) -> tuple[str, dict]:
    """
    URL loader with two-pass scraping:
    Pass 1 → trafilatura (best for articles/papers)
    Pass 2 → LangChain WebBaseLoader (broad fallback)
    """
    meta = {"source": url, "type": "url"}
    text = ""

    # ── Pass 1: trafilatura ──
    try:
        import trafilatura
        downloaded = trafilatura.fetch_url(url)
        if downloaded:
            text = trafilatura.extract(downloaded, include_tables=True, include_comments=False) or ""
    except ImportError:
        pass

    # ── Pass 2: LangChain WebBaseLoader ──
    if len(text) < 300:
        try:
            from langchain_community.document_loaders import WebBaseLoader
            loader = WebBaseLoader(url)
            docs = loader.load()
            text = "\n\n".join(d.page_content for d in docs)
        except Exception as e:
            if not text:
                raise ValueError(f"Could not fetch URL content: {e}")

    # ── Block detection ──
    blocking = ["just a moment", "enable javascript", "checking your browser",
                "cloudflare", "access denied", "please enable javascript"]
    if any(p in text.lower() for p in blocking) or len(text) < 200:
        raise ValueError(
            "⚠️ This website blocks automated scraping (JavaScript/Cloudflare protection).\n\n"
            "✅ Try instead:\n"
            "• Wikipedia: https://en.wikipedia.org/wiki/...\n"
            "• ArXiv HTML: https://arxiv.org/html/...\n"
            "• Semantic Scholar: https://www.semanticscholar.org/\n"
            "• Documentation sites, Medium articles, open blogs"
        )
    
    meta["char_count"] = len(text)
    return text, meta


def load_txt(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """Load plain text, Markdown, or CSV files."""
    meta = {"source": filename, "type": "text"}
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        text = file_bytes.decode("latin-1")
    meta["char_count"] = len(text)
    return text, meta


def load_docx(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """Load Word documents using python-docx."""
    meta = {"source": filename, "type": "docx"}
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    paragraphs.append(row_text)
        text = "\n\n".join(paragraphs)
        meta["char_count"] = len(text)
        return text, meta
    except ImportError:
        raise RuntimeError("python-docx not installed. Run: `pip install python-docx`")


def load_csv(file_bytes: bytes, filename: str) -> tuple[str, dict]:
    """Convert CSV to a readable text representation."""
    meta = {"source": filename, "type": "csv"}
    try:
        import pandas as pd
        df = pd.read_csv(io.BytesIO(file_bytes))
        meta["rows"] = len(df)
        meta["columns"] = list(df.columns)
        text = f"Dataset: {filename}\nShape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
        text += df.to_string(index=False, max_rows=200)
        meta["char_count"] = len(text)
        return text, meta
    except Exception as e:
        raise ValueError(f"CSV parsing failed: {e}")


# ══════════════════════════════════════════════════════════════════════════════
#  TEXT CHUNKING
# ══════════════════════════════════════════════════════════════════════════════
def chunk_text(text: str, metadata: dict) -> list[Document]:
    """
    Industry-standard recursive character chunking with metadata injection.
    
    Chunk size = 1000 chars, overlap = 200 chars:
    - Small enough to fit in LLM context window with many chunks
    - Large enough to preserve meaningful context within a chunk
    - Overlap prevents information loss at chunk boundaries
    """
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
        length_function=len,
    )
    raw_chunks = splitter.split_text(text)
    
    # Inject rich metadata into each chunk for citation-aware retrieval
    docs = []
    for i, chunk in enumerate(raw_chunks):
        doc_meta = {
            **metadata,
            "chunk_index": i,
            "chunk_count": len(raw_chunks),
            "chunk_preview": chunk[:80].replace("\n", " "),
        }
        docs.append(Document(page_content=chunk, metadata=doc_meta))
    
    return docs


# ══════════════════════════════════════════════════════════════════════════════
#  EMBEDDING + VECTOR STORE
# ══════════════════════════════════════════════════════════════════════════════
@st.cache_resource(show_spinner=False)
def get_embeddings():
    """
    Load HuggingFace MiniLM-L6-v2 once and reuse across all sessions.
    Trade-off: local/free vs OpenAI text-embedding-3-small (paid, higher quality).
    """
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )


def build_vector_store(docs: list[Document]) -> FAISS:
    """
    Build FAISS index from documents.
    FAISS = Facebook AI Similarity Search — fast cosine similarity over 384-dim vectors.
    """
    embeddings = get_embeddings()
    vectorstore = FAISS.from_documents(docs, embeddings)
    return vectorstore


def similarity_search(vectorstore: FAISS, query: str, k: int = 4) -> list[Document]:
    """
    Retrieve top-k most semantically relevant chunks.
    k=4 balances context richness vs LLM token limits.
    """
    return vectorstore.similarity_search(query, k=k)


# ══════════════════════════════════════════════════════════════════════════════
#  LLM SETUP
# ══════════════════════════════════════════════════════════════════════════════
def get_llm(api_key: str, temperature: float = 0.2) -> ChatGroq:
    """
    Groq LLaMA-3.3-70B:
    - 70B parameters = strong reasoning for complex academic content
    - Groq inference = faster than OpenAI at same quality tier
    - temperature=0.2 = mostly factual, minimal hallucination
    """
    return ChatGroq(
        api_key=api_key,
        model="llama-3.3-70b-versatile",
        temperature=temperature,
        max_tokens=3000,
    )


# ══════════════════════════════════════════════════════════════════════════════
#  RESEARCH PAPER PROMPT TEMPLATES
# ══════════════════════════════════════════════════════════════════════════════

STRUCTURED_SUMMARY_PROMPT = """You are an expert AI/ML research paper analyst. Analyze the provided context from a research paper and produce a STRUCTURED summary in the EXACT format below. Be thorough, precise, and use technical language appropriately.

Context from the paper:
{context}

Produce this exact structure (use markdown headers exactly as shown):

## 🎯 Problem Statement
[What problem does this paper solve? Why does it matter?]

## 💡 Key Contributions
[List 3-5 specific novel contributions, numbered]

## 🔬 Methodology / Approach
[Explain the method, architecture, or algorithm proposed. Be specific about technical details.]

## 📊 Results & Performance
[State the quantitative results, benchmarks, datasets used, and how they compare to baselines.]

## ⚠️ Limitations
[What are the acknowledged or apparent limitations?]

## 🔭 Future Work
[What directions do the authors suggest? What gaps remain?]

## 🏷️ Domain Tags
[List 4-6 relevant tags: e.g., #TransformerArchitecture, #NLP, #ComputerVision, #Optimization]

If any section's information is NOT in the provided context, write: "[Not found in provided excerpt — try asking a specific question]"
Do not make up information not present in the context."""


CHAT_PROMPT = """You are an expert AI/ML research paper analyst helping a researcher understand a paper. Answer the question accurately and concisely based ONLY on the provided context chunks.

Context (retrieved from the paper):
{context}

Conversation so far:
{history}

User Question: {question}

Rules:
1. Answer ONLY from the context. Do not hallucinate.
2. If the answer is not in the context, say: "This specific information isn't in the retrieved sections. Try rephrasing or asking about a different aspect."
3. Cite section/chunk when helpful (e.g., "According to the methods section...")
4. For mathematical concepts, explain intuitively AND technically.
5. Be conversational but precise.

Answer:"""


QUICK_INSIGHTS_PROMPT = """You are an AI research paper analyst. Read the context and extract the 5 most important insights a researcher would want to know immediately.

Context:
{context}

Format your response EXACTLY as:
**⚡ 5 Key Insights from this Paper**

1. 🔹 **[Insight title]**: [1-2 sentence explanation]
2. 🔹 **[Insight title]**: [1-2 sentence explanation]
3. 🔹 **[Insight title]**: [1-2 sentence explanation]
4. 🔹 **[Insight title]**: [1-2 sentence explanation]
5. 🔹 **[Insight title]**: [1-2 sentence explanation]

**📌 One-line takeaway**: [The single most important thing to remember]

Base everything ONLY on the provided context."""


SECTION_DIVE_PROMPT = """You are a research paper expert. The user wants a DEEP DIVE into a specific section or topic from the paper.

Full retrieved context:
{context}

User's focus area: {section}

Provide:
1. **Overview** — What this section covers (2-3 sentences)
2. **Technical Details** — Explain the core technical content in depth
3. **Why It Matters** — Significance for the field
4. **Connections** — How this relates to other known work (if context supports it)
5. **Key Terms** — Define 3-5 important technical terms used here

Be thorough. Use the context only. If the section isn't covered in the retrieved text, say so."""


def format_context(docs: list[Document]) -> str:
    """Format retrieved Document chunks into a clean context string."""
    parts = []
    for i, doc in enumerate(docs, 1):
        src = doc.metadata.get("source", "unknown")
        chunk_i = doc.metadata.get("chunk_index", "?")
        parts.append(f"[Chunk {chunk_i} | Source: {src}]\n{doc.page_content}")
    return "\n\n---\n\n".join(parts)


def format_history(history: list[dict]) -> str:
    """Format chat history for inclusion in prompt."""
    if not history:
        return "No previous conversation."
    lines = []
    for msg in history[-4:]:  # Only last 4 turns to stay within token limits
        role = "User" if msg["role"] == "user" else "Assistant"
        lines.append(f"{role}: {msg['content'][:300]}")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════════════════
#  LLM INFERENCE FUNCTIONS
# ══════════════════════════════════════════════════════════════════════════════
def run_structured_summary(llm, vectorstore) -> str:
    """Generate a full structured paper summary using broad retrieval."""
    breadth_queries = [
        "problem statement motivation introduction",
        "methodology approach model architecture algorithm",
        "experiments results performance benchmarks",
        "contributions novel innovation",
        "limitations future work conclusion",
    ]
    all_docs = []
    seen_chunks = set()
    for q in breadth_queries:
        docs = similarity_search(vectorstore, q, k=3)
        for d in docs:
            key = d.metadata.get("chunk_index", d.page_content[:50])
            if key not in seen_chunks:
                all_docs.append(d)
                seen_chunks.add(key)
    context = format_context(all_docs[:12])
    prompt = STRUCTURED_SUMMARY_PROMPT.format(context=context)
    response = llm.invoke(prompt)
    return response.content


def run_chat(llm, vectorstore, question: str, history: list) -> tuple[str, list]:
    """Answer a free-form question with conversation history."""
    docs = similarity_search(vectorstore, question, k=4)
    context = format_context(docs)
    hist_str = format_history(history)
    prompt = CHAT_PROMPT.format(context=context, history=hist_str, question=question)
    response = llm.invoke(prompt)
    answer = response.content
    
    # Update history
    history.append({"role": "user", "content": question})
    history.append({"role": "assistant", "content": answer})
    
    return answer, docs, history


def run_quick_insights(llm, vectorstore) -> str:
    """Extract top 5 key insights from the paper."""
    docs = similarity_search(vectorstore, "main findings results key contributions", k=5)
    context = format_context(docs)
    prompt = QUICK_INSIGHTS_PROMPT.format(context=context)
    response = llm.invoke(prompt)
    return response.content


def run_section_dive(llm, vectorstore, section: str) -> str:
    """Deep-dive into a specific section or concept."""
    docs = similarity_search(vectorstore, section, k=5)
    context = format_context(docs)
    prompt = SECTION_DIVE_PROMPT.format(context=context, section=section)
    response = llm.invoke(prompt)
    return response.content


# ══════════════════════════════════════════════════════════════════════════════
#  DOCUMENT PROCESSING PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
def process_document(source_type: str, file_obj=None, url: str = "") -> bool:
    """
    Central pipeline: Load → Chunk → Embed → Store
    Returns True on success, False on failure.
    """
    progress = st.progress(0, text="Initializing...")
    status = st.empty()

    try:
        # ── Step 1: Load ──────────────────────────────────────────────────
        status.markdown('<div class="step-row"><span class="step-icon">📥</span> Loading document...</div>', unsafe_allow_html=True)
        progress.progress(10, text="Loading document...")

        if source_type in ("pdf", "image", "txt", "docx", "csv"):
            file_bytes = file_obj.read()
            filename = file_obj.name
            current_hash = _file_hash(file_bytes)
            
            # Skip reprocessing if same document
            if current_hash == st.session_state.doc_hash:
                progress.empty(); status.empty()
                st.info("⚡ Same document — using cached vector store.")
                return True

            if source_type == "pdf":
                raw_text, meta = load_pdf(file_bytes, filename)
            elif source_type == "image":
                raw_text, meta = load_image(file_bytes, filename)
            elif source_type == "docx":
                raw_text, meta = load_docx(file_bytes, filename)
            elif source_type == "csv":
                raw_text, meta = load_csv(file_bytes, filename)
            else:  # txt, md
                raw_text, meta = load_txt(file_bytes, filename)

        elif source_type == "url":
            current_hash = _file_hash(url.encode())
            if current_hash == st.session_state.doc_hash:
                progress.empty(); status.empty()
                st.info("⚡ Same URL — using cached vector store.")
                return True
            raw_text, meta = load_url(url)
        else:
            raise ValueError(f"Unknown source type: {source_type}")

        progress.progress(30, text="Document loaded ✓")

        # ── Step 2: Chunk ─────────────────────────────────────────────────
        status.markdown('<div class="step-row"><span class="step-icon">✂️</span> Splitting into semantic chunks...</div>', unsafe_allow_html=True)
        progress.progress(45, text="Chunking text...")
        chunks = chunk_text(raw_text, meta)

        progress.progress(55, text=f"Created {len(chunks)} chunks ✓")

        # ── Step 3: Embed + Index ─────────────────────────────────────────
        status.markdown('<div class="step-row"><span class="step-icon">🔢</span> Creating vector embeddings (local MiniLM-L6-v2)...</div>', unsafe_allow_html=True)
        progress.progress(65, text="Embedding chunks...")
        vectorstore = build_vector_store(chunks)
        progress.progress(90, text="Vector index built ✓")

        # ── Step 4: Save to session state ─────────────────────────────────
        st.session_state.vector_store = vectorstore
        st.session_state.doc_hash = current_hash
        st.session_state.doc_metadata = meta
        st.session_state.raw_text = raw_text
        st.session_state.doc_chunks = chunks
        st.session_state.chat_history = []  # Reset chat on new doc
        st.session_state.processing_done = True

        progress.progress(100, text="Ready! ✅")
        time.sleep(0.5)
        progress.empty()
        status.empty()
        return True

    except Exception as e:
        progress.empty()
        status.empty()
        st.error(f"**Processing failed:** {e}")
        return False


# ══════════════════════════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════════════════════════
def render_sidebar() -> tuple[str, bool]:
    """Render sidebar: API key, source selector, file/URL upload. Returns (api_key, ready)."""
    
    with st.sidebar:
        st.markdown("## 🔬 Research Analyzer")
        st.markdown("---")

        # ── API Key ──
        st.markdown("### 🔑 Groq API Key")
        api_key = os.getenv("GROQ_API_KEY", "")
        if api_key:
            st.success("✅ Loaded from `.env`")
        else:
            api_key = st.text_input(
                "Enter Groq API key",
                type="password",
                placeholder="gsk_...",
                help="Free key at https://console.groq.com/keys"
            )
        
        if not api_key:
            st.warning("⚠️ API key required to analyze papers.")
        
        st.markdown("---")

        # ── Source Type ──
        st.markdown("### 📂 Input Source")
        source_options = {
            "📄 PDF (Text-based)": "pdf",
            "🖼️ Image (JPG/JPEG/PNG)": "image",
            "🌐 Website / ArXiv URL": "url",
            "📝 Text / Markdown": "txt",
            "📊 CSV / Data": "csv",
            "📋 Word Document (DOCX)": "docx",
        }
        selected_label = st.selectbox(
            "Select input type",
            list(source_options.keys()),
            index=0,
        )
        source_type = source_options[selected_label]

        # ── File / URL Input ──
        st.markdown("### 📤 Upload / Enter")
        file_obj = None
        url = ""

        if source_type == "url":
            url = st.text_input(
                "Enter URL",
                placeholder="https://arxiv.org/html/...",
            )
            with st.expander("💡 Compatible URLs"):
                st.markdown("""
- **ArXiv HTML**: `https://arxiv.org/html/<id>`
- **Wikipedia**: `https://en.wikipedia.org/wiki/...`
- **Semantic Scholar**: search & paste article URL
- **Medium**: public articles
- **Documentation sites**: docs.python.org, etc.

❌ Avoid: Cloudflare-protected, login-required sites
                """)
        elif source_type == "pdf":
            file_obj = st.file_uploader("Upload PDF", type=["pdf"])
        elif source_type == "image":
            file_obj = st.file_uploader("Upload Image", type=["jpg", "jpeg", "png"])
            with st.expander("🔧 OCR Setup"):
                st.markdown("""
Requires **Tesseract OCR** binary:
```bash
pip install pytesseract pillow
```
Download Tesseract: [GitHub Release](https://github.com/UB-Mannheim/tesseract/wiki)
                """)
        elif source_type == "txt":
            file_obj = st.file_uploader("Upload File", type=["txt", "md", "rst"])
        elif source_type == "csv":
            file_obj = st.file_uploader("Upload CSV", type=["csv"])
        elif source_type == "docx":
            file_obj = st.file_uploader("Upload Word Doc", type=["docx"])
            with st.expander("🔧 Setup"):
                st.markdown("Requires: `pip install python-docx`")

        # ── Process Button ──
        process_ready = (file_obj is not None) or (source_type == "url" and url.strip())
        
        if process_ready:
            if st.button("🚀 Analyze Document", use_container_width=True):
                with st.spinner("Processing..."):
                    success = process_document(source_type, file_obj, url)
                if success:
                    st.success("✅ Ready to analyze!")
                    st.rerun()

        st.markdown("---")

        # ── Document Stats ──
        if st.session_state.processing_done:
            meta = st.session_state.doc_metadata
            st.markdown("### 📊 Document Stats")
            st.markdown(f"""
<div class="metric-box">
    <div class="metric-value">{len(st.session_state.doc_chunks)}</div>
    <div class="metric-label">Chunks Indexed</div>
</div>
""", unsafe_allow_html=True)
            st.markdown("")
            cols = st.columns(2)
            char_count = meta.get("char_count", len(st.session_state.raw_text))
            with cols[0]:
                st.metric("Characters", f"{char_count:,}")
            with cols[1]:
                pages = meta.get("pages", "—")
                st.metric("Pages", pages)
            
            src = meta.get("source", "Unknown")
            doc_type = meta.get("type", "unknown").upper()
            st.markdown(f'<span class="badge badge-blue">{doc_type}</span>', unsafe_allow_html=True)
            if meta.get("ocr_used"):
                st.markdown('<span class="badge badge-amber">OCR</span>', unsafe_allow_html=True)
            st.caption(f"📎 {Path(src).name if source_type != 'url' else src[:40]}")

            # Reset button
            if st.button("🗑️ Clear & Reset", use_container_width=True):
                for k in ["vector_store", "doc_hash", "doc_metadata", "chat_history",
                          "raw_text", "processing_done", "doc_chunks"]:
                    st.session_state[k] = None if k in ["vector_store","doc_hash"] else (
                        [] if "history" in k or "chunks" in k else
                        {} if "metadata" in k else
                        "" if "text" in k else False
                    )
                st.rerun()

        st.markdown("---")
        st.markdown("""
<div style='color:#8b949e;font-size:0.75rem;text-align:center;'>
🔬 RAG Research Analyzer<br>
Groq LLaMA 3.3 · FAISS · MiniLM-L6-v2<br>
LangChain · Streamlit
</div>
""", unsafe_allow_html=True)

    return api_key, st.session_state.processing_done


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN UI TABS
# ══════════════════════════════════════════════════════════════════════════════
def render_main(api_key: str, ready: bool):
    """Render the main area with 4 analysis mode tabs."""

    # ── Header ──
    st.markdown('<div class="app-header">🔬 AI Research Paper Analyzer</div>', unsafe_allow_html=True)
    st.markdown(
        '<div class="sub-header">Industry-Standard RAG Pipeline · '
        'PDF · Images · URLs · DOCX · CSV · TXT</div>',
        unsafe_allow_html=True
    )

    if not ready:
        _render_landing()
        return

    if not api_key:
        st.error("🔑 Please provide your Groq API key in the sidebar.")
        return

    # LLM instance (created fresh per interaction, no caching for API key safety)
    llm = get_llm(api_key)

    # ── Analysis Mode Tabs ──
    tab1, tab2, tab3, tab4 = st.tabs([
        "🔬 Structured Summary",
        "💬 Chat Q&A",
        "⚡ Quick Insights",
        "🔍 Section Deep-Dive"
    ])

    with tab1:
        _render_structured_summary(llm)

    with tab2:
        _render_chat(llm)

    with tab3:
        _render_quick_insights(llm)

    with tab4:
        _render_section_dive(llm)


def _render_landing():
    """Landing page shown when no document is loaded."""
    st.markdown("---")
    cols = st.columns(3)
    
    features = [
        ("📄", "Multi-Format Input", "PDF (text + scanned), Images (OCR), URLs, DOCX, CSV, TXT"),
        ("🧠", "LLaMA 3.3 70B", "State-of-the-art open-source LLM via Groq's ultra-fast inference"),
        ("🔢", "Local Embeddings", "MiniLM-L6-v2 runs on your CPU — no embedding API costs"),
        ("⚡", "FAISS Vector Search", "Millisecond similarity search over document chunks"),
        ("📊", "Structured Analysis", "Auto-extract: Problem, Method, Results, Limitations"),
        ("💬", "Multi-Turn Chat", "Conversation with memory — ask follow-ups naturally"),
    ]
    
    for i, (icon, title, desc) in enumerate(features):
        with cols[i % 3]:
            st.markdown(f"""
<div class="analysis-card">
    <div style="font-size:2rem;margin-bottom:0.5rem">{icon}</div>
    <div style="color:#c9d1d9;font-weight:600;margin-bottom:0.3rem">{title}</div>
    <div style="color:#8b949e;font-size:0.85rem">{desc}</div>
</div>
""", unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("""
<div style='text-align:center;color:#8b949e;padding:2rem;'>
    <div style='font-size:3rem;margin-bottom:1rem'>⬅️</div>
    <div style='font-size:1.1rem;color:#c9d1d9;font-weight:500'>Upload a research paper or paste a URL in the sidebar to begin</div>
    <div style='font-size:0.9rem;margin-top:0.5rem'>Supports PDF, JPG/JPEG/PNG, URLs, TXT, DOCX, CSV</div>
</div>
""", unsafe_allow_html=True)


def _render_structured_summary(llm):
    """Tab 1: Full structured analysis of the paper."""
    st.markdown("### 🔬 Structured Research Paper Summary")
    st.markdown(
        "Automatically extracts **Problem, Contributions, Methodology, Results, "
        "Limitations, Future Work** and domain tags from the paper.",
        help="Uses 5 targeted retrieval queries + LLaMA 3.3 70B for comprehensive coverage."
    )

    if "structured_summary" not in st.session_state:
        st.session_state.structured_summary = ""

    col1, col2 = st.columns([1, 3])
    with col1:
        generate = st.button("🚀 Generate Summary", use_container_width=True)
    with col2:
        if st.session_state.structured_summary:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            download_content = f"# Research Paper Analysis\n\n{st.session_state.structured_summary}"
            st.download_button(
                "⬇️ Download Summary (.md)",
                data=download_content,
                file_name=f"paper_analysis_{timestamp}.md",
                mime="text/markdown",
                use_container_width=True
            )

    if generate:
        with st.spinner("🔍 Retrieving relevant sections and synthesizing analysis..."):
            try:
                summary = run_structured_summary(llm, st.session_state.vector_store)
                st.session_state.structured_summary = summary
            except Exception as e:
                st.error(f"Analysis failed: {e}")

    if st.session_state.structured_summary:
        st.markdown("---")
        st.markdown(st.session_state.structured_summary)


def _render_chat(llm):
    """Tab 2: Multi-turn conversational Q&A."""
    st.markdown("### 💬 Chat with the Paper")
    st.markdown("Ask questions in natural language. Conversation history is maintained across turns.")

    # Display chat history
    history = st.session_state.chat_history
    
    if history:
        for msg in history:
            if msg["role"] == "user":
                st.markdown(f'<div class="chat-user">👤 <b>You:</b> {msg["content"]}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div class="chat-assistant">🤖 <b>Analyst:</b><br>{msg["content"]}</div>', unsafe_allow_html=True)
        
        if st.button("🗑️ Clear Chat", use_container_width=False):
            st.session_state.chat_history = []
            st.rerun()
    else:
        st.markdown("""
<div style='text-align:center;color:#8b949e;padding:2rem;border:1px dashed #30363d;border-radius:8px;margin:1rem 0;'>
    Ask anything about the paper — methodology, results, specific equations, comparisons...
</div>
""", unsafe_allow_html=True)

    # Input box
    st.markdown("---")
    
    # Suggested questions
    if not history:
        st.markdown("**💡 Suggested questions:**")
        suggestions = [
            "What problem does this paper solve?",
            "Explain the proposed method in simple terms",
            "What datasets were used?",
            "How does this compare to previous work?",
        ]
        scols = st.columns(2)
        for i, q in enumerate(suggestions):
            with scols[i % 2]:
                if st.button(f"📌 {q}", use_container_width=True, key=f"sug_{i}"):
                    st.session_state._pending_question = q
                    st.rerun()

    # Handle pending question from suggestion buttons
    pending = st.session_state.pop("_pending_question", None)
    
    question = st.text_area(
        "Your question:",
        value=pending or "",
        placeholder="E.g., What is the main novelty compared to BERT?",
        height=80,
        key="chat_input"
    )

    col1, col2 = st.columns([1, 5])
    with col1:
        ask = st.button("💬 Ask", use_container_width=True)

    if ask and question.strip():
        with st.spinner("🤔 Searching paper and generating answer..."):
            try:
                answer, source_docs, updated_history = run_chat(
                    llm,
                    st.session_state.vector_store,
                    question.strip(),
                    st.session_state.chat_history,
                )
                st.session_state.chat_history = updated_history

                # Show source chunks
                with st.expander(f"📎 Retrieved {len(source_docs)} source chunks", expanded=False):
                    for i, doc in enumerate(source_docs, 1):
                        st.markdown(f"""
<div class="source-box">
    <b>Chunk {i}</b> (index {doc.metadata.get('chunk_index','?')})<br>
    {doc.page_content[:300]}{'...' if len(doc.page_content)>300 else ''}
</div>
""", unsafe_allow_html=True)

                st.rerun()
            except Exception as e:
                st.error(f"Chat failed: {e}")


def _render_quick_insights(llm):
    """Tab 3: 5 bullet-point key insights."""
    st.markdown("### ⚡ Quick Insights")
    st.markdown("Get the **5 most important takeaways** from the paper — ideal for a quick overview before a meeting or review.")

    if "quick_insights" not in st.session_state:
        st.session_state.quick_insights = ""

    if st.button("⚡ Generate Insights", use_container_width=False):
        with st.spinner("🔍 Identifying key findings..."):
            try:
                insights = run_quick_insights(llm, st.session_state.vector_store)
                st.session_state.quick_insights = insights
            except Exception as e:
                st.error(f"Failed: {e}")

    if st.session_state.quick_insights:
        st.markdown("---")
        st.markdown(st.session_state.quick_insights)


def _render_section_dive(llm):
    """Tab 4: Deep-dive into a specific section or concept."""
    st.markdown("### 🔍 Section Deep-Dive")
    st.markdown("Choose a **specific section or concept** to explore in depth.")

    preset_sections = [
        "Abstract and Introduction",
        "Related Work and Background",
        "Methodology and Model Architecture",
        "Experiments and Evaluation",
        "Results and Analysis",
        "Conclusion and Future Work",
        "Loss Functions and Optimization",
        "Dataset Description",
    ]

    col1, col2 = st.columns([2, 1])
    with col1:
        section_input = st.text_input(
            "Section or topic to deep-dive:",
            placeholder="E.g., self-attention mechanism, training procedure, ablation study...",
        )
    with col2:
        preset = st.selectbox("Or pick a preset:", ["— Custom —"] + preset_sections)

    if preset != "— Custom —":
        section_input = preset

    if st.button("🔍 Deep Dive", use_container_width=False) and section_input.strip():
        with st.spinner(f"🔬 Analyzing: {section_input}..."):
            try:
                result = run_section_dive(llm, st.session_state.vector_store, section_input)
                st.markdown("---")
                st.markdown(result)
            except Exception as e:
                st.error(f"Failed: {e}")


# ══════════════════════════════════════════════════════════════════════════════
#  ENTRYPOINT
# ══════════════════════════════════════════════════════════════════════════════
def main():
    api_key, ready = render_sidebar()
    render_main(api_key, ready)


if __name__ == "__main__":
    main()
